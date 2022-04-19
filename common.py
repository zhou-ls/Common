"""
用于封装各种常用函数功能
"""
import logging
import os
import platform
import re
import smtplib
import socket
from collections import Counter
from email.header import Header
from email.mime.text import MIMEText

import qrcode
import requests
from docx import Document
from openpyxl import Workbook
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter, PDFTextExtractionNotAllowed
from pdfminer.pdfparser import PDFParser, PDFDocument


__all__ = ["trying", "remove_space", "get_file_name", "get_host_ip", "count_list", "read_txt_file", "creat_excel",
           "name_repeat", "get_html", "load_data", "bio_sent", "product_ner_train_data", "split_data", "log_print",
           "qr_code", "pdf2word", "send_mail"]


def trying(counts: int):
    """
    装饰器
    传入重试次数，如果失败，重试
    :param counts: 重试次数
    :return:
    """

    def decorator(func):
        def wrapper(*args, **kwargs):
            # nonlocal counts
            for i in range(counts):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    logging.error(e)
                    logging.info('第{}次重试'.format(i + 1))

        return wrapper

    return decorator


def remove_space(text: str):
    """
    移除文本中的空格
    :return:
    """
    return re.sub('\s+', '', text).strip()


def get_file_name(path: str) -> str:
    # 获取不同操作系统(windows, linux)下的路径文件名
    return os.path.basename(path)


def get_host_ip():
    """
    查询本机ip地址和操作系统
    :return: ip
    """
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('8.8.8.8', 80))
        ip = s.getsockname()[0]
    finally:
        s.close()
    return ip, platform.platform()


def count_list(array):
    """
    统计列表中元素出现的频率
    :param array: 列表
    :return: 频率由高到低
    """
    result = Counter(array)
    # 排序
    result = sorted(result.items(), key=lambda x: x[1], reverse=True)
    return result


def read_txt_file(path):
    """
    读取txt的每一行放入列表中
    :param path: txt文件路径
    :return: 列表， txt文件的每一行的内容
    """
    with open(path, 'r', encoding="utf-8") as f:
        content = [_.strip() for _ in f.readlines()]
    return content


def creat_excel():
    """
    :return 当前活动文件及表格
    """
    wb = Workbook()
    wb.create_sheet(index=0, title="Sheet1")
    ac_sheet = wb.active
    return wb, ac_sheet


def name_repeat(drug_list):
    """
    除掉嵌套的药品名称
    :param drug_list: 药品名称列表
    :return: 去重后的药品名称列表
    """
    repeat = []
    drug_list = sorted(drug_list, key=lambda x: len(x), reverse=False)  # 按药品名称长度升序排列
    for m in range(len(drug_list) - 1):
        flag = 0
        for n in range(m + 1, len(drug_list)):
            if drug_list[m] in drug_list[n]:
                flag = 1
        if flag == 0:  # 如果不是子字符串
            repeat.append(drug_list[m])
    repeat.append(drug_list[-1])  # 最后一个字符串长度最长，一定不是子字符串
    return repeat


def get_html(url, headers):
    """
    利用requests爬虫获取网页html
    :param url: 网址路径
    :param headers： 浏览器请求头
    :return 网页html
    """
    n = 0
    try:
        response = requests.get(url, headers=headers, timeout=3)
        response.encoding = response.apparent_encoding
        html = response.text
        return html
    except ConnectionError:
        n += 1
        if n < 4:
            print('连接超时，正在重新爬取数据......')
            return get_html(url, headers)
        else:
            return None


def load_data(filename):
    """
    加载NER数据
    单条格式：[(片段1, 标签1), (片段2, 标签2), (片段3, 标签3), ...]
    """
    D = []
    with open(filename, encoding='utf-8') as f:
        f = f.read()
        for l in f.split('\n\n'):
            if not l:
                continue
            d, last_flag = [], ''
            for c in l.split('\n'):
                char, this_flag = c.split(' ')
                if this_flag == 'O' and last_flag == 'O':
                    d[-1][0] += char
                elif this_flag == 'O' and last_flag != 'O':
                    d.append([char, 'O'])
                elif this_flag[:1] == 'B':
                    d.append([char, this_flag[2:]])
                else:
                    d[-1][0] += char
                last_flag = this_flag
            D.append(d)
    return D


def bio_sent(sent, entity_list, head, tail):
    """
    加工成BIO标注序列, 单个
    Args:
        tail: 标注的实体中间及尾部， string
        head: 标注的实体头， string
        sent: 文本信息
        entity_list: 该文本信息中所包含的所有的实体名称列表,去重后的

    Returns:

    """
    sent = str(sent)
    bio_list = ['O'] * len(sent)
    for entity in entity_list:
        for i in range(0, len(sent) - len(entity) + 1):
            if sent[i:i + len(entity)] == entity:
                bio_list[i] = head
                for j in range(1, len(entity)):
                    bio_list[i + j] = tail
    return sent, bio_list


def product_ner_train_data(result_path, sent_list, entity_list, **kwargs):
    """
    第一步在词典中匹配出该句子中所含有的所有的实体名称，得到该句子对应的实体列表 entity_list
    第二步调用 将 entity_list 中的字符串按照长度由短到长进行排序
    第三步调用 product_ner_train_data() 生成 NER 模型的训练集
    Args:
        result_path:最终形成的NER训练数据存放的地址
        sent_list:文本信息列表
        entity_list:所有的实体名称形成的列表
        **kwargs:  "B-PER,I-PER"=person_list  ...
                    person_list 为某一种实体类型组成的实体列表
                    key为每类实体的标注标签
                    value为每类实体形成的实体列表

    Returns:

    """
    entity_list = sorted(list(set(entity_list)), key=lambda x: len(x), reverse=False)  # 按药品名称长度升序排列
    f = open(result_path, 'w', encoding="utf-8")
    for sent in sent_list:
        sent = str(sent)
        bio_list = ['O'] * len(sent)
        for entity in entity_list:
            for i in range(0, len(sent) - len(entity) + 1):
                if sent[i:i + len(entity)] == entity:  # 以一个实体字符串的长度在该文本信息上形成滑动窗口
                    for key, value in kwargs.items():
                        if entity in value:
                            head, tail = key.split(',')[0], key.split(',')[1]
                            bio_list[i] = head
                            for j in range(1, len(entity)):
                                bio_list[i + j] = tail
        for char, tag in zip(sent, bio_list):
            if not char:
                print(sent)
            f.write(char + ' ' + tag + '\n')
        f.write('\n')


def split_data(result_path, train_path='train.txt', test_path='test.txt', dev_path='dev.txt', train_fold=0.7,
               test_fold=0.2):
    """

    Args:
        result_path: NER数据集
        train_path: 训练数据集
        test_path: 测试数据集
        dev_path: 验证数据集
        train_fold: 训练集所占的比例
        test_fold: 测试集所占的比例

    Returns:训练集、测试集、验证集

    """
    train = open(train_path, 'w', encoding='utf-8')
    test = open(test_path, 'w', encoding='utf-8')
    dev = open(dev_path, 'w', encoding='utf-8')
    content = read_txt_file(result_path)
    for i in range(int(train_fold * len(content))):
        train.write(content[i] + '\n')

    for i in range(int(train_fold * len(content)), int((test_fold + train_fold) * len(content))):
        test.write(content[i] + '\n')

    for i in range(int((test_fold + train_fold) * len(content)), len(content)):
        dev.write(content[i] + '\n')


def log_print(log_path, content):
    """
    打印输出日志
    :param log_path: 日志输出保存路径
    :param content： 保存内容
    :return 网页html
    """
    logging.basicConfig(level=logging.DEBUG,  # 控制台打印的日志级别
                        filename=log_path,
                        filemode='a',  # 模式，有w和a，w就是写模式，每次都会重新写日志，覆盖之前的日志,a是追加模式，默认如果不写的话，就是追加模式
                        format='%(asctime)s - %(levelname)s: %(message)s'  # 日志格式
                        )
    logging.info('\n' + content)


def qr_code(url, path='qrcode.png'):
    # 输入一个链接，生成其对应的二维码
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')
    img.save(path)


def pdf2word(src_path, result_path):
    document = Document()
    # rb以二进制读模式打开本地pdf文件
    fn = open(src_path, 'rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)

    # 提供初始化密码doc.initialize("lianxipython")
    # 如果没有密码 就创建一个空的字符串
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed

    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource, laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource, device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out, "get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    # with open('example_dis.test','a') as f:
                    #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    document.add_paragraph(
                        content, style='ListBullet'  # 添加段落，样式为unordered list类型
                    )
                document.save(result_path)  # 保存这个文档


def send_mail(mail_host,
              mail_user,
              mail_pass,
              sender,
              receiver,
              sender_name,
              receiver_name,
              title,
              content):
    """
    发送邮件
    :param mail_host: 发送邮件的服务器，如 smtp.qq.com
    :param mail_user： 用户  如 123456@qq.com
    :param mail_pass： 授权码
    :param sender： 发送方   如 123456@qq.com
    :param receiver： 接收方 如 123456789@qq.com
    :param sender_name： 发件人名字
    :param receiver_name： 收件人名字
    :param title： 邮件主题
    :param content： 邮件内容
    :return
    """
    message = MIMEText(content, 'html', 'utf-8')  # 邮件内容
    message['From'] = Header(sender_name, 'utf-8')  # 发件人名字
    message['To'] = Header(receiver_name, 'utf-8')  # 收件人名字
    message['Subject'] = Header(title, 'utf-8')
    try:
        smtpObj = smtplib.SMTP()
        smtpObj.connect(mail_host, 25)  # 25 为 SMTP 端口号
        smtpObj.login(mail_user, mail_pass)
        smtpObj.sendmail(sender, receiver, message.as_string())
        print("邮件发送成功")
    except smtplib.SMTPException:
        print("Error: 无法发送邮件")


if __name__ == '__main__':
    print(get_host_ip())
